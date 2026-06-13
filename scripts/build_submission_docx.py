#!/usr/bin/env python3
"""Build submission DOCX with Word equations, figures, and captions.

Features:
- LaTeX → MathML → OMML Word equations
- Display equations with sequential numbering (1), (2), ...
- Figure captions with images
- Tables after References
- Line numbers, page numbers
- Times New Roman 12pt, double spacing
- Submission package in /submission/
"""

from __future__ import annotations

import re
import shutil
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree
from docx.oxml.ns import qn

import latex2mathml.converter

PROJECT = Path(__file__).resolve().parents[1]
DRAFT = PROJECT / "manuscript/draft_v7.md"
FIGURES_DIR = PROJECT / "results/figures_v6"
SUBMISSION_DIR = PROJECT / "submission"
OUTPUT = SUBMISSION_DIR / "Park_2026_WRR_Physical_Credit_Audit.docx"
COPY_DRAFT_NAME = "draft_v7.md"
COPY_PACKAGE_FILES = True
DOC_FONT = 'Times New Roman'
DOC_FONT_SIZE = 12
DOC_FONT_COLOR = RGBColor(0, 0, 0)

# ── OMML conversion ──
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NARY_OPS = {'∑', '∏', '∫', '⋃', '⋂'}
ACCENT_MAP = {'^': '̂', '~': '̃', '¯': '̄', '˙': '̇', '¨': '̈', 'ˇ': '̌'}


def _strip_mml_ns(elem):
    for el in elem.iter():
        if isinstance(el.tag, str) and '}' in el.tag:
            el.tag = el.tag.split('}', 1)[1]


def _get_text(node) -> str:
    parts = []
    if node.text:
        parts.append(node.text)
    for child in node:
        parts.append(_get_text(child))
        if child.tail:
            parts.append(child.tail)
    return ''.join(parts).strip()


def _omml_run(parent_m, text: str):
    r = etree.SubElement(parent_m, f'{{{MATH_NS}}}r')
    rpr = etree.SubElement(r, f'{{{MATH_NS}}}rPr')
    etree.SubElement(rpr, f'{{{MATH_NS}}}sty').set(f'{{{MATH_NS}}}val', 'p')
    t = etree.SubElement(r, f'{{{MATH_NS}}}t')
    t.text = text


def _is_nary_mathml(node) -> bool:
    if node is None:
        return False
    if node.tag in ('msubsup', 'munderover'):
        base = _get_text(node[0]) if len(node) >= 1 else ''
        return bool(base and base in NARY_OPS)
    if node.tag == 'mo':
        base = _get_text(node)
        return bool(base and base in NARY_OPS)
    return False


def _append_nary_from_mathml(mml_node, parent_m):
    base = _get_text(mml_node[0]) if mml_node.tag != 'mo' and len(mml_node) >= 1 else _get_text(mml_node)
    nary = etree.SubElement(parent_m, f'{{{MATH_NS}}}nary')
    narypr = etree.SubElement(nary, f'{{{MATH_NS}}}naryPr')
    chr_el = etree.SubElement(narypr, f'{{{MATH_NS}}}chr')
    chr_el.set(f'{{{MATH_NS}}}val', base)
    limloc = etree.SubElement(narypr, f'{{{MATH_NS}}}limLoc')
    limloc.set(f'{{{MATH_NS}}}val', 'subSup' if mml_node.tag == 'msubsup' else 'undOvr')
    sub = etree.SubElement(nary, f'{{{MATH_NS}}}sub')
    sup = etree.SubElement(nary, f'{{{MATH_NS}}}sup')
    etree.SubElement(nary, f'{{{MATH_NS}}}e')
    if mml_node.tag != 'mo':
        if len(mml_node) >= 2:
            _convert_node(mml_node[1], sub)
        if len(mml_node) >= 3:
            _convert_node(mml_node[2], sup)
    return nary


def _convert_node(mml_node, parent_m):
    tag = mml_node.tag
    if tag in ('math', 'mrow', 'mstyle'):
        children = list(mml_node)
        idx = 0
        while idx < len(children):
            child = children[idx]
            if _is_nary_mathml(child):
                nary = _append_nary_from_mathml(child, parent_m)
                expr = nary.find(f'{{{MATH_NS}}}e')
                idx += 1
                while idx < len(children):
                    next_child = children[idx]
                    if next_child.tag == 'mo' and (next_child.text or '').strip() in ('=', '<', '>', '≤', '≥', '≈', '≠', ',', '.', ';', ':'):
                        break
                    _convert_node(next_child, expr)
                    idx += 1
                continue
            _convert_node(child, parent_m)
            idx += 1
        return

    if tag in ('mi', 'mn', 'mo', 'mtext'):
        text = (mml_node.text or '').strip()
        if text:
            _omml_run(parent_m, text)
        return

    if tag == 'mfrac':
        f = etree.SubElement(parent_m, f'{{{MATH_NS}}}f')
        etree.SubElement(etree.SubElement(f, f'{{{MATH_NS}}}fPr'), f'{{{MATH_NS}}}ctrlPr')
        num = etree.SubElement(f, f'{{{MATH_NS}}}num')
        den = etree.SubElement(f, f'{{{MATH_NS}}}den')
        if len(mml_node) >= 1:
            _convert_node(mml_node[0], num)
        if len(mml_node) >= 2:
            _convert_node(mml_node[1], den)
        return

    if tag == 'msub':
        ssub = etree.SubElement(parent_m, f'{{{MATH_NS}}}sSub')
        base = etree.SubElement(ssub, f'{{{MATH_NS}}}e')
        sub = etree.SubElement(ssub, f'{{{MATH_NS}}}sub')
        if len(mml_node) >= 1:
            _convert_node(mml_node[0], base)
        if len(mml_node) >= 2:
            _convert_node(mml_node[1], sub)
        return

    if tag == 'msup':
        ssup = etree.SubElement(parent_m, f'{{{MATH_NS}}}sSup')
        base = etree.SubElement(ssup, f'{{{MATH_NS}}}e')
        sup = etree.SubElement(ssup, f'{{{MATH_NS}}}sup')
        if len(mml_node) >= 1:
            _convert_node(mml_node[0], base)
        if len(mml_node) >= 2:
            _convert_node(mml_node[1], sup)
        return

    if tag == 'msubsup':
        base = _get_text(mml_node[0]) if len(mml_node) >= 1 else ''
        if base and base in NARY_OPS:
            _append_nary_from_mathml(mml_node, parent_m)
            return
        ssubsup = etree.SubElement(parent_m, f'{{{MATH_NS}}}sSubSup')
        e = etree.SubElement(ssubsup, f'{{{MATH_NS}}}e')
        sub = etree.SubElement(ssubsup, f'{{{MATH_NS}}}sub')
        sup = etree.SubElement(ssubsup, f'{{{MATH_NS}}}sup')
        if len(mml_node) >= 1:
            _convert_node(mml_node[0], e)
        if len(mml_node) >= 2:
            _convert_node(mml_node[1], sub)
        if len(mml_node) >= 3:
            _convert_node(mml_node[2], sup)
        return

    if tag == 'mover':
        accent = _get_text(mml_node[1]) if len(mml_node) >= 2 else ''
        acc = etree.SubElement(parent_m, f'{{{MATH_NS}}}acc')
        accpr = etree.SubElement(acc, f'{{{MATH_NS}}}accPr')
        chr_el = etree.SubElement(accpr, f'{{{MATH_NS}}}chr')
        chr_el.set(f'{{{MATH_NS}}}val', ACCENT_MAP.get(accent, accent or '̂'))
        e = etree.SubElement(acc, f'{{{MATH_NS}}}e')
        if len(mml_node) >= 1:
            _convert_node(mml_node[0], e)
        return

    if tag in ('munderover', 'munder', 'mover'):
        for child in mml_node:
            _convert_node(child, parent_m)
        return

    if tag == 'msqrt':
        rad = etree.SubElement(parent_m, f'{{{MATH_NS}}}rad')
        radpr = etree.SubElement(rad, f'{{{MATH_NS}}}radPr')
        deg_hide = etree.SubElement(radpr, f'{{{MATH_NS}}}degHide')
        deg_hide.set(f'{{{MATH_NS}}}val', '1')
        etree.SubElement(rad, f'{{{MATH_NS}}}deg')
        e = etree.SubElement(rad, f'{{{MATH_NS}}}e')
        for child in mml_node:
            _convert_node(child, e)
        return

    if tag == 'mfenced':
        d = etree.SubElement(parent_m, f'{{{MATH_NS}}}d')
        dpr = etree.SubElement(d, f'{{{MATH_NS}}}dPr')
        beg = etree.SubElement(dpr, f'{{{MATH_NS}}}begChr')
        beg.set(f'{{{MATH_NS}}}val', mml_node.get('open', '('))
        end = etree.SubElement(dpr, f'{{{MATH_NS}}}endChr')
        end.set(f'{{{MATH_NS}}}val', mml_node.get('close', ')'))
        e = etree.SubElement(d, f'{{{MATH_NS}}}e')
        for child in mml_node:
            _convert_node(child, e)
        return

    text = _get_text(mml_node)
    if text:
        _omml_run(parent_m, text)


def latex_to_omml(latex_str: str) -> etree._Element | None:
    """Convert LaTeX to Word OMML without leaking visible raw syntax."""
    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mml = etree.fromstring(mathml_str.encode('utf-8'))
        _strip_mml_ns(mml)
        omath = etree.Element(f'{{{MATH_NS}}}oMath')
        _convert_node(mml, omath)
        return omath
    except Exception:
        return None


def add_inline_equation(paragraph, latex_str: str):
    """Add an inline OMML equation to a paragraph."""
    omml = latex_to_omml(latex_str)
    if omml is not None:
        paragraph._element.append(omml)
    else:
        # Fallback stays plain so failed conversion is visible without changing style.
        run = paragraph.add_run(latex_str)
        set_run_format(run)


def add_display_equation(doc, latex_str: str, eq_num: int):
    """Add a display equation with number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    omml = latex_to_omml(latex_str)
    if omml is not None:
        # Wrap in oMathPara for display
        omathpara = etree.Element(qn('m:oMathPara'))
        omathpara.append(omml)
        p._element.append(omathpara)
    else:
        run = p.add_run(latex_str)
        set_run_format(run)

    # Add equation number
    run = p.add_run(f'   ({eq_num})')
    set_run_format(run)
    return p


def extract_figure_captions(markdown_text: str) -> dict[int, str]:
    """Extract figure captions from the draft's Figure Captions section."""
    captions: dict[int, str] = {}
    in_caption_section = False
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped == "## Figure Captions":
            in_caption_section = True
            continue
        if in_caption_section and stripped.startswith("## "):
            break
        if not in_caption_section:
            continue
        match = re.match(r"^\*\*Figure\s+(\d+)\.\*\*\s*(.*)$", stripped)
        if match:
            captions[int(match.group(1))] = match.group(2).strip()
    return captions


def extract_title(markdown_text: str) -> str:
    """Return the first Markdown H1 title."""
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            return stripped[2:].strip()
    return "Untitled Manuscript"


def extract_h2_section(markdown_text: str, section_name: str) -> list[str]:
    """Return non-empty lines from a named H2 section."""
    lines: list[str] = []
    in_section = False
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped == f"## {section_name}":
            in_section = True
            continue
        if in_section and stripped.startswith("## "):
            break
        if in_section and stripped:
            lines.append(stripped)
    return lines


def strip_ordered_marker(text: str) -> str:
    """Remove a leading Markdown ordered-list marker if present."""
    return re.sub(r"^\d+\.\s+", "", text.strip())


def set_run_format(run, *, bold: bool = False, italic: bool | None = None, size: int = DOC_FONT_SIZE):
    run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.name = DOC_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = DOC_FONT_COLOR


def set_paragraph_format(paragraph):
    paragraph.style = 'Normal'
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


# ── Document builder ──
def build_docx():
    doc = Document()
    text = DRAFT.read_text()

    # Style setup
    style = doc.styles['Normal']
    style.font.name = DOC_FONT
    style.font.size = Pt(DOC_FONT_SIZE)
    style.font.color.rgb = DOC_FONT_COLOR
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Line numbers
        sectPr = section._sectPr
        ln_num = etree.SubElement(sectPr, qn('w:lnNumType'))
        ln_num.set(qn('w:countBy'), '1')
        ln_num.set(qn('w:restart'), 'continuous')
        # Page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        set_run_format(run)
        fld_begin = etree.SubElement(run._element, qn('w:fldChar'))
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run2 = para.add_run()
        instr = etree.SubElement(run2._element, qn('w:instrText'))
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run3 = para.add_run()
        fld_end = etree.SubElement(run3._element, qn('w:fldChar'))
        fld_end.set(qn('w:fldCharType'), 'end')

    # ── Helper ──
    def add_heading(text, level=1):
        p = doc.add_paragraph()
        set_paragraph_format(p)
        run = p.add_run(text)
        set_run_format(run)
        return p

    def add_runs_with_inline_math(paragraph, text, font_size=Pt(12)):
        parts = re.split(r'(\\\(.+?\\\)|\$[^$]+\$)', text)
        for part in parts:
            if part.startswith('$') and part.endswith('$'):
                latex = part[1:-1]
                add_inline_equation(paragraph, latex)
            elif part.startswith(r'\(') and part.endswith(r'\)'):
                latex = part[2:-2]
                add_inline_equation(paragraph, latex)
            else:
                if part:
                    run = paragraph.add_run(part)
                    set_run_format(run, size=int(font_size.pt if hasattr(font_size, 'pt') else DOC_FONT_SIZE))

    def add_body(text):
        if not text.strip():
            return None
        p = doc.add_paragraph()
        set_paragraph_format(p)
        add_runs_with_inline_math(p, text.strip(), Pt(12))
        return p

    def add_reference(text):
        p = doc.add_paragraph()
        set_paragraph_format(p)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        parts = re.split(r'(\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
            else:
                run = p.add_run(part)
            set_run_format(run)

    # ── Title page ──
    add_heading(extract_title(text), level=0)

    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run("Eungyu Park")
    set_run_format(run)

    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run("Department of Geology, Kyungpook National University, Daegu 41566, South Korea")
    set_run_format(run)

    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run("Corresponding author: Eungyu Park (egpark@knu.ac.kr)")
    set_run_format(run)

    # ── Parse main text ──
    lines = text.split('\n')

    eq_counter = [0]  # mutable for closure
    tables_collected = []
    in_references = False
    current_ref = ""
    skip_sections = {'Key Points', 'Plain Language Summary', 'Keywords',
                     'Data Availability Statement', 'Acknowledgments',
                     'Declaration of Competing Interests', 'Author Contributions',
                     'Tables', 'Figure Captions'}
    skip_until_next_h2 = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip title (already added)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            i += 1
            continue

        # H2 sections
        if stripped.startswith('## '):
            section_title = stripped[3:].strip()

            # Check if should skip
            if any(section_title.startswith(s) for s in skip_sections):
                skip_until_next_h2 = True
                # But collect tables
                if section_title == 'Tables':
                    # Collect ALL tables in the Tables section at once
                    i += 1
                    current_title = None
                    current_lines = []
                    while i < len(lines):
                        s = lines[i].strip()
                        if s.startswith('## ') and not s.startswith('### '):
                            break  # End of Tables section
                        if s.startswith('### Table'):
                            # Save previous table if any
                            if current_title is not None:
                                tables_collected.append((current_title, current_lines))
                            current_title = s[4:]  # Remove "### "
                            current_lines = []
                        elif s.startswith('|'):
                            current_lines.append(lines[i])
                        i += 1
                    # Save last table
                    if current_title is not None:
                        tables_collected.append((current_title, current_lines))
                    continue
                i += 1
                continue

            skip_until_next_h2 = False

            if section_title == 'References':
                in_references = True
                p = doc.add_paragraph()
                p.paragraph_format.page_break_before = True
                add_heading('References', level=1)
                i += 1
                continue

            if section_title == 'Figure Captions':
                skip_until_next_h2 = True
                i += 1
                continue

            # Numbered sections
            m = re.match(r'(\d+)\.\s*(.*)', section_title)
            if m:
                add_heading(f'{m.group(1)}. {m.group(2)}', level=1)
            else:
                add_heading(section_title, level=1)
            i += 1
            continue

        if skip_until_next_h2:
            if stripped.startswith('## '):
                skip_until_next_h2 = False
            else:
                i += 1
                continue

        # H3 sections
        if stripped.startswith('### '):
            sub = stripped[4:].strip()
            add_heading(sub, level=2)
            i += 1
            continue

        # References
        if in_references:
            if stripped.startswith('## '):
                if current_ref:
                    add_reference(current_ref)
                    current_ref = ""
                in_references = False
                continue
            if stripped and stripped[0].isupper():
                if current_ref:
                    add_reference(current_ref)
                current_ref = stripped
            elif stripped:
                current_ref += ' ' + stripped
            elif current_ref:
                add_reference(current_ref)
                current_ref = ""
            i += 1
            continue

        # Table lines → placeholder
        if stripped.startswith('|'):
            t_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                t_lines.append(lines[i])
                i += 1
            tables_collected.append((f"Table {len(tables_collected)+1}", t_lines))
            p = doc.add_paragraph()
            set_paragraph_format(p)
            run = p.add_run(f'[Insert Table {len(tables_collected)} about here]')
            set_run_format(run)
            continue

        # Check for display equations: $$...$$ or multi-line $$...$$ (N) patterns
        if stripped.startswith('$$'):
            if not stripped.endswith('$$') and '$$' not in stripped[2:]:
                equation_lines = [stripped[2:]]
                i += 1
                closing_number = None
                while i < len(lines):
                    candidate = lines[i].strip()
                    closing_match = re.match(r'^(.*?)\$\$\s*\((\d+)\)\s*$', candidate)
                    if closing_match:
                        equation_lines.append(closing_match.group(1).strip())
                        closing_number = int(closing_match.group(2))
                        break
                    if candidate.endswith('$$'):
                        equation_lines.append(candidate[:-2].strip())
                        break
                    equation_lines.append(candidate)
                    i += 1
                latex = ' '.join(part for part in equation_lines if part).strip()
                if closing_number is not None:
                    eq_counter[0] = max(eq_counter[0], closing_number)
                    add_display_equation(doc, latex, closing_number)
                else:
                    eq_counter[0] += 1
                    add_display_equation(doc, latex, eq_counter[0])
                i += 1
                continue
            # Try to match $$...$$  (N) pattern
            display_match = re.match(r'^\$\$(.*?)\$\$\s*\((\d+)\)\s*$', stripped)
            if display_match:
                latex = display_match.group(1).strip()
                eq_num = int(display_match.group(2))
                eq_counter[0] = max(eq_counter[0], eq_num)
                add_display_equation(doc, latex, eq_num)
                i += 1
                continue
            # Try pure $$...$$ without number
            elif stripped.endswith('$$'):
                latex = stripped[2:-2].strip()
                eq_counter[0] += 1
                add_display_equation(doc, latex, eq_counter[0])
                i += 1
                continue

        # Normal paragraph
        if stripped:
            # Check for figure caption patterns
            if stripped.startswith('**Figure ') or stripped.startswith('**Fig'):
                clean = stripped.replace('**', '')
                p = doc.add_paragraph()
                set_paragraph_format(p)
                m = re.match(r'(Figure \d+\.)\s*(.*)', clean)
                if m:
                    run = p.add_run(m.group(1) + ' ')
                    set_run_format(run)
                    run = p.add_run(m.group(2))
                    set_run_format(run)
                else:
                    run = p.add_run(clean)
                    set_run_format(run)
            else:
                add_body(stripped)

        i += 1

    # Flush last reference
    if current_ref:
        add_reference(current_ref)

    # ── Data Availability ──
    add_heading("Data Availability Statement", level=1)
    for line in extract_h2_section(text, "Data Availability Statement"):
        add_body(line)

    # ── Acknowledgments ──
    acknowledgments = extract_h2_section(text, "Acknowledgments")
    if acknowledgments:
        add_heading("Acknowledgments", level=1)
        for line in acknowledgments:
            add_body(line)

    # ── Competing Interests ──
    add_heading("Declaration of Competing Interests", level=1)
    for line in extract_h2_section(text, "Declaration of Competing Interests"):
        add_body(line)

    # ── Author Contributions ──
    add_heading("Author Contributions", level=1)
    for line in extract_h2_section(text, "Author Contributions"):
        add_body(line)

    # ── Tables (after References) ──
    for title, t_lines in tables_collected:
        p = doc.add_paragraph()
        set_paragraph_format(p)
        p.paragraph_format.page_break_before = True
        run = p.add_run(title)
        set_run_format(run)

        if t_lines:
            rows = []
            for tl in t_lines:
                if tl.strip().startswith('|') and not re.match(r'^\|[\s\-:]+\|', tl.strip()):
                    cells = [c.strip().replace('**', '') for c in tl.strip().split('|')[1:-1]]
                    if cells:
                        rows.append(cells)
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < ncols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            for p2 in cell.paragraphs:
                                set_paragraph_format(p2)
                                for run in p2.runs:
                                    set_run_format(run)

    # ── Figure Captions page ──
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    add_heading("Figure Captions", level=1)

    captions = extract_figure_captions(text)

    for fig_num in sorted(captions):
        p = doc.add_paragraph()
        set_paragraph_format(p)
        run = p.add_run(f'Figure {fig_num}. ')
        set_run_format(run)
        add_runs_with_inline_math(p, captions.get(fig_num, ''), Pt(12))

    # ── Figures with images ──
    fig_names = {
        1: 'fig01_study_area_map.png',
        2: 'fig02_falsification_heatmap.png',
        3: 'fig03_architecture_diversity.png',
        4: 'fig04_lambda_sensitivity.png',
        5: 'fig05_nuisance_parameter_collapse.png',
        6: 'fig06_physical_credit_ladder.png',
        7: 'fig07_clean_canonical80_central_summary.png',
        8: 'fig08_clean_canonical80_regime_paired.png',
    }

    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    add_heading("Figures", level=1)

    for fig_num in sorted(fig_names):
        fig_path = FIGURES_DIR / fig_names[fig_num]
        if fig_path.exists():
            p = doc.add_paragraph()
            p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(str(fig_path), width=Inches(5.5))
            except Exception:
                run = p.add_run(f'[Figure {fig_num} image]')

            # Caption below figure
            cap = doc.add_paragraph()
            set_paragraph_format(cap)
            run = cap.add_run(f'Figure {fig_num}. ')
            set_run_format(run)
            add_runs_with_inline_math(cap, captions.get(fig_num, ''), Pt(12))

    # ── Save ──
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))

    if COPY_PACKAGE_FILES:
        # Copy figures to submission
        for fig_name in fig_names.values():
            src = FIGURES_DIR / fig_name
            if src.exists():
                shutil.copy2(src, SUBMISSION_DIR / fig_name)

        # Copy markdown draft
        shutil.copy2(DRAFT, SUBMISSION_DIR / COPY_DRAFT_NAME)

    print(f"DOCX saved: {OUTPUT}")
    print(f"Submission package: {SUBMISSION_DIR}")
    print(f"Files in submission/:")
    for f in sorted(SUBMISSION_DIR.iterdir()):
        print(f"  {f.name} ({f.stat().st_size // 1024} KB)")
    print(f"\nEquations added: {eq_counter[0]}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build submission DOCX from a manuscript Markdown source.")
    parser.add_argument("--draft", type=Path, default=DRAFT)
    parser.add_argument("--figures-dir", type=Path, default=FIGURES_DIR)
    parser.add_argument("--submission-dir", type=Path, default=SUBMISSION_DIR)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--copy-draft-name", default=COPY_DRAFT_NAME)
    parser.add_argument("--skip-package-copy", action="store_true")
    args = parser.parse_args()

    DRAFT = args.draft
    FIGURES_DIR = args.figures_dir
    SUBMISSION_DIR = args.submission_dir
    OUTPUT = args.output
    COPY_DRAFT_NAME = args.copy_draft_name
    COPY_PACKAGE_FILES = not args.skip_package_copy

    build_docx()
